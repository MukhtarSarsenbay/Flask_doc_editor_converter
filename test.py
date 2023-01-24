from flask import Flask, request, render_template, send_file
from docx import Document
from fpdf import FPDF

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def modify_and_download():
    if request.method == 'POST':
        # Get the user's input
        text_to_replace = request.form['text_to_replace']
        replaced_text = request.form['replaced_text']

        # Open the .docx file
        doc = Document('original.docx')

        # Make changes to the .docx file
        for para in doc.paragraphs:
            if text_to_replace in para.text:
                para.text = replaced_text

        # Save the modified .docx file
        doc.save('modified.docx')

        # Convert the .docx file to a .pdf file
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.cell(200, 10, txt=para.text, ln=1, align="C")
        pdf.output("modified.pdf").encode('latin1')

        # Make both files available for download
        return send_file('modified.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         attachment_filename='modified.docx'), send_file('modified.pdf',
                         mimetype='application/pdf',
                         as_attachment=True,
                         attachment_filename='modified.pdf')
    return render_template('index.html')
@app.route('/download_docx')
def download_docx():
    return send_file('modified.docx', as_attachment=True)

@app.route('/download_pdf')
def download_pdf():
    return send_file('modified.pdf', as_attachment=True)
if __name__ == '_main_':
    app.run(debug=True)